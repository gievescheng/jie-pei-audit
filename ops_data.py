from __future__ import annotations

import csv
import html
import io
import json
import re
from copy import deepcopy
from datetime import date, datetime
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
from werkzeug.utils import secure_filename

from runtime_paths import PRIVATE_DATA_DIR


BASE_DIR = Path(__file__).resolve().parent
_STORAGE_ROOT = (PRIVATE_DATA_DIR / "operations").resolve()
_UPLOAD_ROOT = (_STORAGE_ROOT / "uploads").resolve()

DEFAULT_NONCONFORMANCES = [
    {
        "id": "NC-2025-001",
        "date": "2025-11-12",
        "dept": "品管課",
        "type": "人員作業",
        "description": "在品檢室進行作業時，不慎撞到已洗完待檢測的玻璃，導致掉落地板破裂。",
        "severity": "輕微",
        "rootCause": "品管作業員在作業時不慎撞到 FOSB 盒，導致盒子掉落並摔破玻璃片。",
        "correctiveAction": "作業員立即清理現場並確認無玻璃碎片殘留，另行包裝破碎玻璃片至包裝袋並隔離儲存；主管提醒作業員注意各物品擺放。",
        "responsible": "林佑翰",
        "dueDate": "",
        "status": "已關閉",
        "closeDate": "2025-11-12",
        "effectiveness": "有效",
        "source_file": "",
    },
    {
        "id": "NC-2026-001",
        "date": "2026-01-28",
        "dept": "品檢課",
        "type": "人員作業",
        "description": "品檢人員於 AOI 檢驗後取出 NG 品放置 FOSB，中途玻璃脫落導致破片。",
        "severity": "輕微",
        "rootCause": "AOI 測試判定 NG 後，取出放置 NG FOSB 時操作不慎，玻璃脫落破片。",
        "correctiveAction": "加強作業員操作訓練，明確規範 FOSB 取放動作要領，並更新作業 SOP。",
        "responsible": "朱姿霖",
        "dueDate": "",
        "status": "處理中",
        "closeDate": "",
        "effectiveness": "",
        "source_file": "",
    },
]

DEFAULT_AUDIT_PLANS = [
    {
        "id": "IA-2025-01",
        "year": 2025,
        "period": "下半年",
        "scheduledDate": "2025-09-04",
        "dept": "全廠",
        "scope": "MP-01,MP-02,MP-03,MP-04,MP-05,MP-06",
        "auditor": "蔡有為",
        "auditee": "程鼎智",
        "status": "已完成",
        "actualDate": "2025-09-05",
        "findings": 4,
        "ncCount": 1,
        "source_file": "",
        "attachment_paths": [
            "9 內部稽核管理程序/記錄/內部稽核114年度/9.5 品質稽核報告書.docx",
            "9 內部稽核管理程序/記錄/內部稽核114年度/9.3 內部稽核查檢表.xlsx",
        ],
    },
    {
        "id": "IA-2025-02",
        "year": 2025,
        "period": "上半年",
        "scheduledDate": "2025-06-23",
        "dept": "品管課",
        "scope": "MP-11,MP-13",
        "auditor": "蔡有為",
        "auditee": "程鼎智",
        "status": "已完成",
        "actualDate": "2025-06-23",
        "findings": 1,
        "ncCount": 1,
        "source_file": "",
        "attachment_paths": [
            "9 內部稽核管理程序/記錄/內部稽核114年度/9.4內部稽核矯正通知單.docx",
            "9 內部稽核管理程序/表單/9.1 內部稽核計畫表.DOCx",
        ],
    },
]

KIND_META = {
    "nonconformance": {"file": "nonconformances.json", "defaults": DEFAULT_NONCONFORMANCES, "allowed": {".docx", ".xlsx", ".pdf"}},
    "auditplan": {"file": "audit_plans.json", "defaults": DEFAULT_AUDIT_PLANS, "allowed": {".docx", ".xlsx", ".pdf"}},
    "environment": {"file": "environment_records.json", "defaults": [], "allowed": {".xlsx", ".csv"}},
}


def set_storage_root(root: Path | str) -> None:
    global _STORAGE_ROOT, _UPLOAD_ROOT
    _STORAGE_ROOT = Path(root).resolve()
    _UPLOAD_ROOT = (_STORAGE_ROOT / "uploads").resolve()
    _ensure_dirs()


def _ensure_dirs() -> None:
    _STORAGE_ROOT.mkdir(parents=True, exist_ok=True)
    _UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)
    for kind in KIND_META:
        (_UPLOAD_ROOT / kind).mkdir(parents=True, exist_ok=True)


def _store_path(kind: str) -> Path:
    _ensure_dirs()
    return _STORAGE_ROOT / KIND_META[kind]["file"]


def _now_iso() -> str:
    return datetime.now().replace(microsecond=0).isoformat()


def _read_json(path: Path) -> list[dict]:
    if not path.exists():
        return []
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return []


def _write_json(path: Path, payload: list[dict]) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _parse_date(value) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    text = str(value).strip()
    if not text:
        return ""
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y-%m-%d %H:%M:%S", "%Y/%m/%d %H:%M:%S", "%d/%m/%Y", "%d/%m/%Y %H:%M:%S"):
        try:
            return datetime.strptime(text, fmt).date().isoformat()
        except ValueError:
            pass
    match = re.search(r"(\d{2,4})[/-](\d{1,2})[/-](\d{1,2})", text)
    if not match:
        return ""
    year = int(match.group(1))
    if year < 1911:
        year += 1911
    try:
        return date(year, int(match.group(2)), int(match.group(3))).isoformat()
    except ValueError:
        return ""


def _coerce_int(value, default: int = 0) -> int:
    try:
        return int(float(str(value).replace(",", "").strip()))
    except Exception:
        return default


def _coerce_float(value, default: float = 0.0) -> float:
    try:
        return float(str(value).replace(",", "").strip())
    except Exception:
        return default


def _coerce_optional_float(value):
    if value is None:
        return ""
    text = str(value).strip()
    if not text:
        return ""
    return _coerce_float(text, 0.0)


def _parse_datetime_value(value):
    if value is None:
        return None
    if isinstance(value, datetime):
        return value
    text = str(value).strip()
    if not text:
        return None
    for fmt in (
        "%Y-%m-%d %H:%M:%S",
        "%Y/%m/%d %H:%M:%S",
        "%d/%m/%Y %H:%M:%S",
        "%d/%m/%Y %H:%M",
        "%Y-%m-%d %H:%M",
        "%Y/%m/%d %H:%M",
    ):
        try:
            return datetime.strptime(text, fmt)
        except ValueError:
            pass
    return None


def _format_datetime_iso(value) -> str:
    parsed = _parse_datetime_value(value)
    return parsed.replace(microsecond=0).isoformat() if parsed else ""


def compute_environment_result(record: dict) -> str:
    p05 = _coerce_int(record.get("particles05"))
    p5 = _coerce_int(record.get("particles5"))
    temp = _coerce_float(record.get("temp"))
    humidity = _coerce_float(record.get("humidity"))
    pressure = _coerce_float(record.get("pressure"), 999.0)
    if p05 > 1000 or p5 > 35 or (temp and (temp > 23 or temp < 21)) or (humidity and (humidity > 50 or humidity < 40)) or (pressure != 999.0 and pressure < 10):
        return "不合格"
    if p05 > 800 or p5 > 20 or (temp and temp > 22.5) or (humidity and humidity > 48):
        return "警告"
    return "合格"


def _generate_id(kind: str, record: dict, items: list[dict]) -> str:
    existing_ids = {item.get("id", "") for item in items}
    if kind == "nonconformance":
        year = (_parse_date(record.get("date")) or date.today().isoformat())[:4]
        prefix = f"NC-{year}-"
        width = 3
    elif kind == "auditplan":
        year = (_parse_date(record.get("scheduledDate")) or date.today().isoformat())[:4]
        prefix = f"IA-{year}-"
        width = 2
    else:
        prefix = "ENV-"
        width = 3
    seq = 1
    for item in items:
        item_id = str(item.get("id", ""))
        if item_id.startswith(prefix):
            try:
                seq = max(seq, int(item_id.split("-")[-1]) + 1)
            except Exception:
                pass
    candidate = f"{prefix}{seq:0{width}d}"
    while candidate in existing_ids:
        seq += 1
        candidate = f"{prefix}{seq:0{width}d}"
    return candidate


def _normalize_nonconformance(record: dict, items: list[dict]) -> dict:
    normalized = {
        "id": str(record.get("id") or "").strip(),
        "date": _parse_date(record.get("date")),
        "dept": str(record.get("dept") or "").strip(),
        "type": str(record.get("type") or "製程異常").strip(),
        "description": str(record.get("description") or "").strip(),
        "severity": str(record.get("severity") or "輕微").strip(),
        "rootCause": str(record.get("rootCause") or "").strip(),
        "correctiveAction": str(record.get("correctiveAction") or "").strip(),
        "responsible": str(record.get("responsible") or "").strip(),
        "dueDate": _parse_date(record.get("dueDate")),
        "status": str(record.get("status") or "待處理").strip(),
        "closeDate": _parse_date(record.get("closeDate")),
        "effectiveness": str(record.get("effectiveness") or "").strip(),
        "source_file": str(record.get("source_file") or "").strip(),
    }
    if not normalized["id"]:
        normalized["id"] = _generate_id("nonconformance", normalized, items)
    return normalized


def _normalize_auditplan(record: dict, items: list[dict]) -> dict:
    scheduled = _parse_date(record.get("scheduledDate"))
    actual = _parse_date(record.get("actualDate"))
    year = _coerce_int(record.get("year"), 0) or (int(scheduled[:4]) if scheduled else date.today().year)
    period = str(record.get("period") or "").strip() or ("上半年" if scheduled and int(scheduled[5:7]) <= 6 else "下半年")
    normalized = {
        "id": str(record.get("id") or "").strip(),
        "year": year,
        "period": period,
        "scheduledDate": scheduled,
        "dept": str(record.get("dept") or "").strip(),
        "scope": str(record.get("scope") or "").strip(),
        "auditor": str(record.get("auditor") or "").strip(),
        "auditee": str(record.get("auditee") or "").strip(),
        "status": str(record.get("status") or ("已完成" if actual else "計畫中")).strip(),
        "actualDate": actual,
        "findings": _coerce_int(record.get("findings"), 0),
        "ncCount": _coerce_int(record.get("ncCount"), 0),
        "source_file": str(record.get("source_file") or "").strip(),
        "attachment_paths": [str(item).strip() for item in (record.get("attachment_paths") or []) if str(item).strip()],
    }
    if normalized["source_file"] and normalized["source_file"] not in normalized["attachment_paths"]:
        normalized["attachment_paths"].insert(0, normalized["source_file"])
    if not normalized["id"]:
        normalized["id"] = _generate_id("auditplan", normalized, items)
    return normalized


def _normalize_environment(record: dict, items: list[dict]) -> dict:
    measured_at = _format_datetime_iso(record.get("measuredAt") or record.get("dateTime") or record.get("datetime"))
    normalized_date = _parse_date(record.get("date"))
    if not normalized_date and measured_at:
        normalized_date = measured_at[:10]
    normalized = {
        "id": str(record.get("id") or "").strip(),
        "date": normalized_date,
        "measuredAt": measured_at,
        "location": str(record.get("location") or "潔淨室A區").strip(),
        "point": str(record.get("point") or "").strip(),
        "particles03": _coerce_int(record.get("particles03"), 0),
        "particles05": _coerce_int(record.get("particles05"), 0),
        "particles1": _coerce_int(record.get("particles1"), 0),
        "particles5": _coerce_int(record.get("particles5"), 0),
        "temp": _coerce_optional_float(record.get("temp")),
        "humidity": _coerce_optional_float(record.get("humidity")),
        "pressure": _coerce_optional_float(record.get("pressure")),
        "operator": str(record.get("operator") or "").strip(),
        "result": str(record.get("result") or "").strip(),
        "source_file": str(record.get("source_file") or "").strip(),
    }
    if not normalized["id"]:
        normalized["id"] = _generate_id("environment", normalized, items)
    if normalized["point"] and (not normalized["location"] or normalized["location"] == "潔淨室A區"):
        normalized["location"] = f"粒子計數點 {normalized['point']}"
    if not normalized["result"]:
        normalized["result"] = compute_environment_result(normalized)
    return normalized


def normalize_record(kind: str, record: dict, items: list[dict]) -> dict:
    if kind == "nonconformance":
        normalized = _normalize_nonconformance(record, items)
    elif kind == "auditplan":
        normalized = _normalize_auditplan(record, items)
    else:
        normalized = _normalize_environment(record, items)
    existing = next((item for item in items if item.get("id") == normalized["id"]), None)
    normalized["created_at"] = str(record.get("created_at") or (existing or {}).get("created_at") or _now_iso())
    normalized["updated_at"] = _now_iso()
    return normalized


def load_records(kind: str) -> list[dict]:
    raw = _read_json(_store_path(kind))
    if not raw:
        defaults = []
        for item in deepcopy(KIND_META[kind]["defaults"]):
            defaults.append(normalize_record(kind, item, defaults))
        _write_json(_store_path(kind), defaults)
        return defaults
    normalized = []
    for item in raw:
        normalized.append(normalize_record(kind, item, normalized))
    _write_json(_store_path(kind), normalized)
    return normalized


def save_records(kind: str, items: list[dict]) -> list[dict]:
    normalized = []
    for item in items:
        normalized.append(normalize_record(kind, item, normalized))
    _write_json(_store_path(kind), normalized)
    return normalized


def upsert_records(kind: str, records: list[dict], replace_source_file: str = "") -> tuple[list[dict], list[dict]]:
    items = load_records(kind)
    replace_source_file = str(replace_source_file or "").strip()
    if replace_source_file:
        items = [item for item in items if str(item.get("source_file") or "").strip() != replace_source_file]
    saved = []
    for record in records:
        current = dict(record)
        if replace_source_file and str(current.get("source_file") or "").strip() == replace_source_file:
            current.pop("id", None)
        normalized = normalize_record(kind, current, items)
        index = next((i for i, item in enumerate(items) if item.get("id") == normalized["id"]), None)
        if index is None:
            items.append(normalized)
        else:
            normalized["created_at"] = items[index].get("created_at", normalized["created_at"])
            items[index] = normalized
        saved.append(normalized)
    return save_records(kind, items), saved


def delete_record(kind: str, record_id: str) -> tuple[list[dict], bool]:
    items = load_records(kind)
    kept = [item for item in items if item.get("id") != record_id]
    deleted = len(kept) != len(items)
    if deleted:
        kept = save_records(kind, kept)
    return (kept if deleted else items), deleted


def filter_environment_records(start: str = "", end: str = "") -> list[dict]:
    records = load_records("environment")
    start_date = _parse_date(start)
    end_date = _parse_date(end)
    return [
        item for item in records
        if (not start_date or not item.get("date") or item["date"] >= start_date)
        and (not end_date or not item.get("date") or item["date"] <= end_date)
    ]


def delete_environment_range(start: str = "", end: str = "") -> tuple[list[dict], int]:
    records = load_records("environment")
    start_date = _parse_date(start)
    end_date = _parse_date(end)
    kept = []
    removed = 0
    for item in records:
        rec_date = item.get("date", "")
        in_range = False
        if rec_date:
            if start_date and end_date and start_date <= rec_date <= end_date:
                in_range = True
            elif start_date and not end_date and rec_date >= start_date:
                in_range = True
            elif end_date and not start_date and rec_date <= end_date:
                in_range = True
        if in_range:
            removed += 1
        else:
            kept.append(item)
    if start_date or end_date:
        kept = save_records("environment", kept)
    return kept, removed


def summarize_environment(records: list[dict]) -> dict:
    return {
        "total": len(records),
        "passed": sum(1 for item in records if item.get("result") == "合格"),
        "warning": sum(1 for item in records if item.get("result") == "警告"),
        "failed": sum(1 for item in records if item.get("result") == "不合格"),
    }


def save_uploaded_file(kind: str, file_storage) -> tuple[Path, str]:
    ext = Path(file_storage.filename or "").suffix.lower()
    if ext not in KIND_META[kind]["allowed"]:
        raise ValueError(f"Unsupported file type: {ext or 'unknown'}")
    _ensure_dirs()
    target_dir = _UPLOAD_ROOT / kind
    filename = secure_filename(file_storage.filename or "") or f"upload{ext or '.bin'}"
    target = target_dir / f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
    file_storage.save(target)
    return target, f"uploads/{kind}/{target.name}"


def _clean_extracted_value(value) -> str:
    if value is None:
        return ""
    text = str(value).replace("\xa0", " ").replace("\r", "\n")
    text = text.replace("\n", " / ")
    text = re.sub(r"\s+", " ", text)
    return text.strip(" |:/?\t")


def _flatten_docx(path: Path) -> str:
    doc = Document(path)
    lines = [_clean_extracted_value(p.text) for p in doc.paragraphs if _clean_extracted_value(p.text)]
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cleaned = _clean_extracted_value(cell.text)
                if cleaned and (not cells or cells[-1] != cleaned):
                    cells.append(cleaned)
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _docx_row_values(path: Path) -> list[list[str]]:
    doc = Document(path)
    rows = []
    for table in doc.tables:
        for row in table.rows:
            values = []
            for cell in row.cells:
                cleaned = _clean_extracted_value(cell.text)
                if cleaned and (not values or values[-1] != cleaned):
                    values.append(cleaned)
            if values:
                rows.append(values)
    return rows


def _extract_inline_value(text: str, labels: list[str], stop_labels: list[str] | None = None) -> str:
    if not text:
        return ""
    label_group = "|".join(re.escape(label) for label in labels)
    if stop_labels:
        stop_group = "|".join(re.escape(label) for label in stop_labels)
        pattern = rf"(?:{label_group})\s*[：:]?\s*(.+?)(?=(?:{stop_group})\s*[：:]?|$)"
    else:
        pattern = rf"(?:{label_group})\s*[：:]?\s*(.+)"
    match = re.search(pattern, text, flags=re.IGNORECASE | re.DOTALL)
    return _clean_extracted_value(match.group(1)).lstrip("?: ") if match else ""


def _pick_first(*values) -> str:
    for value in values:
        cleaned = _clean_extracted_value(value)
        if cleaned:
            return cleaned
    return ""


def _extract_dates(text: str) -> list[str]:
    values = []
    for match in re.finditer(r"(\d{2,4}[./-]\d{1,2}[./-]\d{1,2})", text or ""):
        parsed = _parse_date(match.group(1).replace('.', '/'))
        if parsed and parsed not in values:
            values.append(parsed)
    return values


def _flatten_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(text)
    return "\n".join(parts)


def _flatten_xlsx(path: Path) -> str:
    wb = load_workbook(path, data_only=True, read_only=True)
    lines = []
    try:
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = [str(value).strip() for value in row if value not in (None, "")]
                if cells:
                    lines.append(" | ".join(cells))
    finally:
        wb.close()
    return "\n".join(lines)


def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        return _flatten_docx(path)
    if ext == ".pdf":
        return _flatten_pdf(path)
    if ext == ".xlsx":
        return _flatten_xlsx(path)
    if ext == ".csv":
        return path.read_text(encoding="utf-8-sig", errors="ignore")
    return ""


def _extract_value(text: str, labels: list[str]) -> str:
    if not text:
        return ""
    label_group = "|".join(re.escape(label) for label in labels)
    stop_group = "|".join(
        re.escape(token)
        for token in [
            "編號", "日期", "發現單位", "發現部門", "不符合事項說明", "問題描述", "原因分析", "根本原因",
            "矯正措施", "預防措施", "預定完成日期", "到期日期", "責任人", "負責人", "結案日期",
            "稽核部門", "受稽部門", "稽核範圍", "稽核員", "受稽人", "受稽對象", "狀態",
            "實際日期", "發現數", "發現項數", "不符合數", "NC數",
        ]
    )
    match = re.search(
        rf"(?:{label_group})\s*[:：]?\s*(.+?)(?=(?:{stop_group})\s*[:：]|\n|$)",
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    return re.sub(r"\s+", " ", match.group(1).strip()) if match else ""


def _first_date(text: str) -> str:
    match = re.search(r"(\d{2,4}[/-]\d{1,2}[/-]\d{1,2}(?:\s+\d{1,2}:\d{2}:\d{2})?)", text)
    return _parse_date(match.group(1)) if match else ""


def _map_headers(row: list[str], aliases: dict[str, list[str]]) -> dict[str, int]:
    mapping = {}
    lowered = [str(cell or "").strip().lower() for cell in row]
    for field, tokens in aliases.items():
        normalized_tokens = [token.lower() for token in tokens]
        exact_match = next((idx for idx, header in enumerate(lowered) if header in normalized_tokens), None)
        if exact_match is not None:
            mapping[field] = exact_match
            continue
        for idx, header in enumerate(lowered):
            if any(token.lower() in header for token in tokens):
                mapping[field] = idx
                break
    return mapping


def _iter_rows(path: Path) -> list[list]:
    if path.suffix.lower() == ".csv":
        text = path.read_text(encoding="utf-8-sig", errors="ignore")
        return [list(row) for row in csv.reader(io.StringIO(text))]
    wb = load_workbook(path, data_only=True, read_only=True)
    try:
        return [list(row) for ws in wb.worksheets for row in ws.iter_rows(values_only=True)]
    finally:
        wb.close()


def _iter_sheet_rows(path: Path) -> list[list[list]]:
    if path.suffix.lower() == ".csv":
        return [_iter_rows(path)]
    wb = load_workbook(path, data_only=True, read_only=True)
    try:
        return [[list(row) for row in ws.iter_rows(values_only=True)] for ws in wb.worksheets]
    finally:
        wb.close()


def _parse_nonconformance_docx(path: Path, stored_ref: str) -> dict | None:
    rows = _docx_row_values(path)
    if not rows:
        return None
    text = _extract_text(path)
    row1 = rows[0] if rows else []
    row2 = rows[1] if len(rows) > 1 else []
    row3 = rows[2] if len(rows) > 2 else []
    row4 = rows[3] if len(rows) > 3 else []
    row5 = rows[4] if len(rows) > 4 else []
    row6 = rows[5] if len(rows) > 5 else []

    description_line = " ".join(row2)
    root_line = " ".join(row3)
    corrective_line = " ".join(row4)
    preventive_line = " ".join(row5)
    close_line = " ".join(row6)

    dates = _extract_dates(text)
    due_dates = _extract_dates(corrective_line) + _extract_dates(preventive_line)
    id_match = re.search(r"(?:編號|報告編號)\s*[：: ]*([A-Za-z0-9()\-]+)", text[:240])
    dept_cell = row1[1] if len(row1) > 1 and str(row1[0]).startswith("發現單位") else (row1[0] if row1 and not str(row1[0]).endswith("：") else "")

    draft = {
        "id": id_match.group(1) if id_match else _extract_inline_value(text, ["編號", "報告編號"]),
        "date": _pick_first(row1[-1] if row1 else "", _extract_inline_value(text, ["不符合日期", "發現日期", "日期"]), dates[0] if dates else ""),
        "dept": _pick_first(dept_cell, _extract_inline_value(text, ["發現單位", "發現部門", "部門"])),
        "type": "製程異常",
        "description": _pick_first(_extract_inline_value(description_line, ["不符合事項說明"], ["不符合單位主管", "發現者/稽核員", "發現者", "稽核員"]), _extract_inline_value(text, ["不符合事項說明", "問題描述", "不符合事項"])),
        "severity": "重大" if "重大" in text else "輕微",
        "rootCause": _pick_first(_extract_inline_value(root_line, ["原因分析", "根本原因", "異常原因調查分析", "Root Cause Analysis"]), _extract_inline_value(text, ["原因分析", "根本原因", "Root Cause Analysis"])),
        "correctiveAction": _pick_first(_extract_inline_value(corrective_line, ["矯正措施"], ["預定完成日期"]), _extract_inline_value(preventive_line, ["防止再發措施", "預防措施"], ["預定完成日期"]), _extract_inline_value(text, ["矯正措施", "改善措施", "防止再發措施"])),
        "responsible": _pick_first(_extract_inline_value(description_line, ["發現者/稽核員", "發現者", "稽核員"]), _extract_inline_value(text, ["責任人", "負責人", "發現者/稽核員"])),
        "dueDate": due_dates[-1] if due_dates else "",
        "status": "已關閉" if ("結案日期" in close_line and _extract_dates(close_line)) else ("處理中" if text.strip() else "待處理"),
        "closeDate": _pick_first(_extract_dates(close_line)[0] if _extract_dates(close_line) else "", dates[-1] if len(dates) > 1 and "■" in close_line else ""),
        "effectiveness": "有效" if "有效" in text or "結案" in close_line else "",
        "source_file": stored_ref,
    }
    return draft


def parse_nonconformance_import(path: Path, stored_ref: str) -> dict:
    text = _extract_text(path)
    draft = _parse_nonconformance_docx(path, stored_ref) if path.suffix.lower() == ".docx" else None
    if draft is None:
        type_map = {
            "文件": "文件不符",
            "客訴": "客戶投訴",
            "客戶": "客戶投訴",
            "量測": "量測異常",
            "檢驗": "量測異常",
            "來料": "來料不合格",
            "人員": "人員作業",
        }
        nc_type = next((value for keyword, value in type_map.items() if keyword in text), "製程異常")
        draft = {
            "id": _extract_value(text, ["編號", "報告編號"]),
            "date": _extract_value(text, ["發現日期", "不符合日期", "日期"]) or _first_date(text),
            "dept": _extract_value(text, ["發現單位", "發現部門", "部門"]),
            "type": nc_type,
            "description": _extract_value(text, ["不符合事項說明", "問題描述", "不符合事項"]),
            "severity": "重大" if "重大" in text else "輕微",
            "rootCause": _extract_value(text, ["原因分析", "根本原因"]),
            "correctiveAction": _extract_value(text, ["矯正措施", "改善措施"]),
            "responsible": _extract_value(text, ["責任人", "負責人", "發現者/稽核員"]),
            "dueDate": _extract_value(text, ["預定完成日期", "到期日"]),
            "status": "已關閉" if _extract_value(text, ["結案日期"]) else ("處理中" if text.strip() else "待處理"),
            "closeDate": _extract_value(text, ["結案日期"]),
            "effectiveness": "有效" if "有效" in text else "",
            "source_file": stored_ref,
        }
    preview = normalize_record("nonconformance", draft, load_records("nonconformance"))
    missing_fields = [field for field in ("date", "dept", "description", "responsible") if not preview.get(field)]
    return {"draft": preview, "missing_fields": missing_fields, "source_file": stored_ref}


def _parse_auditplan_docx(path: Path, stored_ref: str) -> list[dict]:
    doc = Document(path)
    if not doc.tables:
        return []
    rows = _docx_row_values(path)
    if len(rows) < 2:
        return []
    schedule_text = " ".join(rows[0])
    scheduled = _pick_first(_extract_inline_value(schedule_text, ["稽核時間", "稽核日期"]), _first_date(schedule_text))
    scheduled = _parse_date(scheduled)
    existing = load_records("auditplan")
    parsed = []
    for row in rows[2:]:
        if len(row) >= 5:
            dept = _pick_first(row[1])
            auditee = _pick_first(row[2])
            auditor = _pick_first(row[3])
            scope = _pick_first(row[4])
        else:
            padded = row + [""] * (6 - len(row))
            dept = _pick_first(padded[1], padded[2])
            auditee = _pick_first(padded[3])
            auditor = _pick_first(padded[4])
            scope = _pick_first(padded[5])
        if not any([dept, auditee, auditor, scope]):
            continue
        record = {
            "scheduledDate": scheduled,
            "dept": dept,
            "scope": scope,
            "auditor": auditor,
            "auditee": auditee,
            "status": "計畫中",
            "source_file": stored_ref,
            "attachment_paths": [stored_ref],
        }
        parsed.append(normalize_record("auditplan", record, existing + parsed))
    return parsed


def parse_auditplan_import(path: Path, stored_ref: str) -> dict:
    rows = _iter_rows(path) if path.suffix.lower() == ".xlsx" else []
    aliases = {
        "id": ["稽核編號", "編號", "audit id"],
        "year": ["年度", "year"],
        "period": ["期間", "期別", "period"],
        "scheduledDate": ["預定日期", "計畫日期", "scheduled", "日期"],
        "dept": ["稽核部門", "受稽部門", "受稽核單位", "department", "dept"],
        "scope": ["稽核範圍", "稽核內容", "scope"],
        "auditor": ["稽核員", "auditor"],
        "auditee": ["受稽人", "受稽對象", "受稽核人員", "auditee"],
        "status": ["狀態", "status"],
        "actualDate": ["實際日期", "完成日期", "actual"],
        "findings": ["發現數", "發現項數", "findings"],
        "ncCount": ["不符合數", "nc數", "nccount"],
    }
    records = []
    if rows:
        header_idx = None
        mapping = {}
        for idx, row in enumerate(rows[:20]):
            candidate = _map_headers([str(cell or "") for cell in row], aliases)
            if "scheduledDate" in candidate and ("dept" in candidate or "scope" in candidate):
                header_idx = idx
                mapping = candidate
                break
        if header_idx is not None:
            existing = load_records("auditplan")
            for row in rows[header_idx + 1 :]:
                if all(cell in (None, "") for cell in row):
                    continue
                record = {field: row[col] if col < len(row) else "" for field, col in mapping.items()}
                record["source_file"] = stored_ref
                record["attachment_paths"] = [stored_ref]
                if any(str(record.get(field) or "").strip() for field in ("dept", "auditor", "scheduledDate", "scope", "auditee")):
                    records.append(normalize_record("auditplan", record, existing + records))
    if not records and path.suffix.lower() == ".docx":
        records = _parse_auditplan_docx(path, stored_ref)
    if not records:
        text_only = _extract_text(path)
        actual = _extract_value(text_only, ["實際日期", "完成日期"])
        draft = {
            "id": _extract_value(text_only, ["稽核編號", "編號"]),
            "year": "",
            "period": "",
            "scheduledDate": _extract_value(text_only, ["預定日期", "稽核日期", "計畫日期", "日期"]) or _first_date(text_only),
            "dept": _extract_value(text_only, ["稽核部門", "受稽部門", "受稽核單位", "部門"]),
            "scope": _extract_value(text_only, ["稽核範圍", "稽核內容", "範圍"]),
            "auditor": _extract_value(text_only, ["稽核員"]),
            "auditee": _extract_value(text_only, ["受稽人", "受稽對象", "受稽核人員"]),
            "status": "已完成" if actual else "計畫中",
            "actualDate": actual,
            "findings": _coerce_int(_extract_value(text_only, ["發現數", "發現項數"]), 0),
            "ncCount": _coerce_int(_extract_value(text_only, ["不符合數", "NC數"]), 0),
            "source_file": stored_ref,
            "attachment_paths": [stored_ref],
        }
        records = [normalize_record("auditplan", draft, load_records("auditplan"))]
    preview = [{**record, "missing_fields": [field for field in ("scheduledDate", "dept", "auditor", "scope") if not record.get(field)]} for record in records]
    return {"records": preview, "source_file": stored_ref}


def _parse_environment_tsi(path: Path, stored_ref: str) -> list[dict]:
    rows = _iter_rows(path)
    header_idx = next((idx for idx, row in enumerate(rows[:40]) if "Record Number" in [str(cell or "").strip() for cell in row] and "Date & Time" in [str(cell or "").strip() for cell in row]), None)
    if header_idx is None or header_idx + 1 >= len(rows):
        return []
    header_row = rows[header_idx]
    sample_row = rows[header_idx + 1]
    date_col = next((idx for idx, value in enumerate(header_row) if str(value or "").strip() == "Date & Time"), None)
    location_col = next((idx for idx, value in enumerate(header_row) if str(value or "").strip() == "Location Number"), None)
    record_col = next((idx for idx, value in enumerate(header_row) if str(value or "").strip() == "Record Number"), None)
    p05_col = p1_col = p5_col = None
    for idx, value in enumerate(sample_row[:-1]):
        try:
            size = float(value)
            float(sample_row[idx + 1])
        except Exception:
            continue
        if abs(size - 0.5) < 0.001:
            p05_col = idx + 1
        elif abs(size - 1.0) < 0.001:
            p1_col = idx + 1
        elif abs(size - 5.0) < 0.001:
            p5_col = idx + 1
    existing = load_records("environment")
    parsed = []
    for row in rows[header_idx + 1 :]:
        if record_col is None or date_col is None or record_col >= len(row) or date_col >= len(row) or not row[record_col] or not row[date_col]:
            continue
        record = {
            "date": _parse_date(row[date_col]),
            "measuredAt": _format_datetime_iso(row[date_col]),
            "point": str(row[location_col]) if location_col is not None and location_col < len(row) and row[location_col] not in (None, "") else "",
            "location": f"粒子計數點 {row[location_col]}" if location_col is not None and location_col < len(row) and row[location_col] not in (None, "") else "潔淨室A區",
            "particles03": 0,
            "particles05": row[p05_col] if p05_col is not None and p05_col < len(row) else 0,
            "particles1": row[p1_col] if p1_col is not None and p1_col < len(row) else 0,
            "particles5": row[p5_col] if p5_col is not None and p5_col < len(row) else 0,
            "temp": "",
            "humidity": "",
            "pressure": "",
            "operator": "",
            "source_file": stored_ref,
        }
        parsed.append(normalize_record("environment", record, existing + parsed))
    return parsed


def _parse_environment_clean_long(path: Path, stored_ref: str) -> list[dict]:
    wb = load_workbook(path, data_only=True, read_only=True)
    try:
        worksheet = wb["Clean_Long"] if "Clean_Long" in wb.sheetnames else None
        if worksheet is None:
            return []
        rows = list(worksheet.iter_rows(values_only=True))
    finally:
        wb.close()
    header_idx = next(
        (
            idx
            for idx, row in enumerate(rows[:20])
            if [str(cell or "").strip() for cell in row[:4]] == ["Date", "Point", "DateTime", "Record"]
        ),
        None,
    )
    if header_idx is None:
        return []
    existing = load_records("environment")
    parsed = []
    compact_point = 0
    compact_last_dt = None
    for row in rows[header_idx + 1 :]:
        if not row or row[0] in (None, ""):
            continue
        channel_map = {}
        point_value = row[1] if len(row) > 1 else None
        timestamp = row[2] if len(row) > 2 else None
        date_value = row[0]
        channel_pairs = ((4, 5), (6, 7), (8, 9))
        point_label = ""
        if point_value in (None, "") and timestamp in (None, ""):
            channel_pairs = ((3, 5), (7, 8), (11, 12))
            compact_dt = _parse_datetime_value(date_value)
            if compact_dt is not None:
                if (
                    compact_last_dt is None
                    or compact_dt.date() != compact_last_dt.date()
                    or (compact_dt - compact_last_dt).total_seconds() > 30 * 60
                    or compact_point >= 14
                ):
                    compact_point = 1
                else:
                    compact_point += 1
                compact_last_dt = compact_dt
                point_label = str(compact_point)
                location = f"粒子計數點 {point_label} ({compact_dt.strftime('%H:%M:%S')})"
            else:
                point_label = ""
                location = "粒子計數器"
            measured_at = _format_datetime_iso(compact_dt or date_value)
        else:
            point_label = str(point_value)
            location = f"粒子計數點 {point_value}"
            measured_at = _format_datetime_iso(timestamp or date_value)
            if isinstance(timestamp, datetime):
                location = f"{location} ({timestamp.strftime('%H:%M:%S')})"
            elif timestamp not in (None, ""):
                time_match = re.search(r"(\d{1,2}:\d{2}:\d{2})", str(timestamp))
                if time_match:
                    location = f"{location} ({time_match.group(1)})"
        for size_idx, count_idx in channel_pairs:
            if size_idx >= len(row) or count_idx >= len(row):
                continue
            try:
                size = float(row[size_idx])
                count_value = row[count_idx]
                if count_value in (None, ""):
                    continue
            except Exception:
                continue
            channel_map[round(size, 1)] = _coerce_int(count_value, 0)
        record = {
            "date": _parse_date(date_value),
            "measuredAt": measured_at,
            "location": location,
            "point": point_label,
            "particles03": channel_map.get(0.3, 0),
            "particles05": channel_map.get(0.5, 0),
            "particles1": channel_map.get(1.0, 0),
            "particles5": channel_map.get(5.0, 0),
            "temp": "",
            "humidity": "",
            "pressure": "",
            "operator": "",
            "source_file": stored_ref,
        }
        if record["date"]:
            parsed.append(normalize_record("environment", record, existing + parsed))
    return parsed


def parse_environment_import(path: Path, stored_ref: str) -> dict:
    rows = _iter_rows(path)
    aliases = {
        "date": ["??", "date"],
        "measuredAt": ["datetime", "date & time", "date time", "measured at", "timestamp", "time"],
        "point": ["point", "location number"],
        "location": ["??", "??", "location"],
        "particles03": ["0.3", "0.3?m", "0.3um", ">=0.3"],
        "temp": ["??", "temp"],
        "humidity": ["??", "humidity", "rh"],
        "pressure": ["??", "??", "pressure", "pa"],
        "particles05": ["0.5", "0.5?m", "0.5um", ">=0.5"],
        "particles1": ["1?m", "1um", ">=1"],
        "particles5": ["5.0", "5.0um", "5?m", ">=5"],
        "operator": ["???", "????", "operator"],
        "result": ["??", "??", "result"],
    }
    records = []
    if path.suffix.lower() == ".xlsx":
        records = _parse_environment_clean_long(path, stored_ref)
    if not records and path.suffix.lower() == ".xlsx":
        records = _parse_environment_tsi(path, stored_ref)
    if not records:
        for sheet_rows in _iter_sheet_rows(path):
            header_idx = None
            mapping = {}
            for idx, row in enumerate(sheet_rows[:30]):
                candidate = _map_headers([str(cell or "") for cell in row], aliases)
                if ("date" in candidate or "measuredAt" in candidate) and ("particles05" in candidate or "temp" in candidate or "humidity" in candidate):
                    header_idx = idx
                    mapping = candidate
                    break
            if header_idx is None:
                continue
            existing = load_records("environment")
            for row in sheet_rows[header_idx + 1 :]:
                if all(cell in (None, "") for cell in row):
                    continue
                record = {field: row[col] if col < len(row) else "" for field, col in mapping.items()}
                record["source_file"] = stored_ref
                if not record.get("date") and record.get("measuredAt"):
                    record["date"] = record.get("measuredAt")
                if not _parse_date(record.get("date")):
                    continue
                records.append(normalize_record("environment", record, existing + records))
            if records:
                break
    preview = [{**record, "missing_fields": [field for field in ("date", "location") if not record.get(field)]} for record in records]
    return {"records": preview, "summary": summarize_environment(preview), "source_file": stored_ref}


def parse_import(kind: str, file_storage) -> dict:
    saved_path, stored_ref = save_uploaded_file(kind, file_storage)
    if kind == "nonconformance":
        return parse_nonconformance_import(saved_path, stored_ref)
    if kind == "auditplan":
        return parse_auditplan_import(saved_path, stored_ref)
    return parse_environment_import(saved_path, stored_ref)


def _resolve_storage_path(stored_path: str) -> Path | None:
    cleaned = str(stored_path or "").strip().replace("\\", "/").lstrip("/")
    if not cleaned or cleaned.startswith("."):
        return None
    if cleaned.startswith("uploads/"):
        candidate = (_STORAGE_ROOT / cleaned).resolve()
        root = _STORAGE_ROOT.resolve()
    else:
        candidate = (BASE_DIR / cleaned).resolve()
        root = BASE_DIR.resolve()
    try:
        candidate.relative_to(root)
    except ValueError:
        return None
    if not candidate.exists() or not candidate.is_file() or candidate.name.startswith("."):
        return None
    if candidate.suffix.lower() not in {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".csv", ".txt", ".ppt", ".pptx"}:
        return None
    return candidate


def get_serving_path(stored_path: str) -> Path | None:
    return _resolve_storage_path(stored_path)


def build_text_preview_html(stored_path: str) -> str | None:
    file_path = get_serving_path(stored_path)
    if file_path is None:
        return None
    text = _extract_text(file_path)
    if not text.strip():
        return None
    safe_text = html.escape(text[:120000])
    return f"""<!doctype html>
<html lang="zh-Hant">
<head>
  <meta charset="utf-8" />
  <title>{html.escape(file_path.name)}</title>
  <style>
    body {{ font-family: 'Microsoft JhengHei', 'Noto Sans TC', sans-serif; background:#f8fafc; color:#0f172a; margin:0; }}
    header {{ padding:12px 16px; background:#0f172a; color:#e2e8f0; font-size:14px; font-weight:700; }}
    pre {{ white-space:pre-wrap; word-break:break-word; margin:0; padding:16px; font-family:'Microsoft JhengHei', 'Noto Sans TC', monospace; font-size:13px; line-height:1.65; }}
  </style>
</head>
<body>
  <header>{html.escape(file_path.name)}</header>
  <pre>{safe_text}</pre>
</body>
</html>"""


def list_auditplan_attachments(record_id: str) -> list[dict]:
    record = next((item for item in load_records("auditplan") if item.get("id") == record_id), None)
    if record is None:
        raise KeyError(record_id)
    seen = set()
    attachments = []
    for stored_path in record.get("attachment_paths") or []:
        if not stored_path or stored_path in seen:
            continue
        seen.add(stored_path)
        file_path = get_serving_path(stored_path)
        if file_path is None:
            attachments.append({"name": Path(stored_path).name, "path": stored_path, "exists": False, "previewable": False, "text_previewable": False})
            continue
        suffix = file_path.suffix.lower()
        text_previewable = suffix in {".docx", ".xlsx", ".csv", ".txt"}
        attachments.append(
            {
                "name": file_path.name,
                "path": stored_path,
                "exists": True,
                "previewable": suffix == ".pdf",
                "text_previewable": text_previewable,
                "view_url": f"/api/files/view?path={stored_path}",
                "download_url": f"/api/files/download?path={stored_path}",
                "preview_text_url": f"/api/files/preview-text?path={stored_path}" if text_previewable else "",
            }
        )
    return attachments
