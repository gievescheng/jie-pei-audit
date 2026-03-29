from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document as WordDocument
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


COMPANY_NAME = '潔沛企業有限公司'
REPORT_CODE = 'AA-V2-CMP-01'
HEADER_FILL = PatternFill('solid', fgColor='1F4E78')
SECTION_FILL = PatternFill('solid', fgColor='D9EAF7')
ALERT_FILL = PatternFill('solid', fgColor='FDECEC')
GOOD_FILL = PatternFill('solid', fgColor='E8F5E9')


def _apply_header(row):
    for cell in row:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(vertical='center', horizontal='center')


def _fit_columns(ws):
    widths = {}
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is None:
                continue
            text = str(cell.value)
            widths[cell.column] = max(widths.get(cell.column, 0), min(len(text) + 2, 60))
    for col_idx, width in widths.items():
        ws.column_dimensions[get_column_letter(col_idx)].width = width


def _set_cell_shading(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def _set_table_header(row, fill_hex='1F4E78'):
    for cell in row.cells:
        _set_cell_shading(cell, fill_hex)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = None
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('第 ')
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)
    paragraph.add_run(' 頁')


def _add_key_value_table(doc: WordDocument, pairs: list[tuple[str, str]]):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for key, value in pairs:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = value
        _set_cell_shading(row.cells[0], 'D9EAF7')
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()


def _add_list_table(doc: WordDocument, title: str, headers: list[str], rows: list[list[str]]):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    _set_table_header(table.rows[0])
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            row.cells[idx].text = value
    doc.add_paragraph()


def build_document_compare_workbook(compare_data: dict) -> bytes:
    wb = Workbook()
    summary_ws = wb.active
    summary_ws.title = '摘要'

    left_doc = compare_data.get('left_document', {})
    right_doc = compare_data.get('right_document', {})
    left_only = compare_data.get('left_only_issues', [])
    right_only = compare_data.get('right_only_issues', [])
    added_lines = compare_data.get('added_lines', [])
    removed_lines = compare_data.get('removed_lines', [])
    citations = compare_data.get('citations', [])

    summary_ws['A1'] = '文件差異比對報告'
    summary_ws['A1'].font = Font(size=16, bold=True)
    summary_ws['A3'] = '產出時間'
    summary_ws['B3'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    summary_ws['A4'] = '左側文件'
    summary_ws['B4'] = left_doc.get('title', '')
    summary_ws['A5'] = '右側文件'
    summary_ws['B5'] = right_doc.get('title', '')
    summary_ws['A6'] = '左側來源'
    summary_ws['B6'] = left_doc.get('source_path', '')
    summary_ws['A7'] = '右側來源'
    summary_ws['B7'] = right_doc.get('source_path', '')
    summary_ws['A8'] = '文字相似度'
    summary_ws['B8'] = compare_data.get('similarity', '')
    summary_ws['A9'] = 'Prompt 版本'
    summary_ws['B9'] = compare_data.get('prompt_version', '')
    summary_ws['A10'] = '需人工覆核'
    summary_ws['B10'] = '是' if compare_data.get('needs_human_review') else '否'
    summary_ws['A11'] = '版次結論'
    summary_ws['B11'] = compare_data.get('version_change_conclusion', '')

    for row_idx in range(3, 12):
        summary_ws[f'A{row_idx}'].font = Font(bold=True)
        summary_ws[f'A{row_idx}'].fill = SECTION_FILL

    summary_ws['B11'].alignment = Alignment(wrap_text=True, vertical='top')

    summary_ws['A13'] = '比對摘要'
    summary_ws['A13'].font = Font(bold=True)
    summary_ws['A13'].fill = SECTION_FILL
    summary_ws['A14'] = compare_data.get('summary', '')
    summary_ws['A14'].alignment = Alignment(wrap_text=True, vertical='top')

    summary_ws['A16'] = '差異統計'
    summary_ws['A16'].font = Font(bold=True)
    summary_ws['A16'].fill = SECTION_FILL
    summary_ws.append(['項目', '數量'])
    _apply_header(summary_ws[17])
    summary_ws.append(['右側新增內容', len(added_lines)])
    summary_ws.append(['左側移除內容', len(removed_lines)])
    summary_ws.append(['僅左側存在的規則缺口', len(left_only)])
    summary_ws.append(['僅右側存在的規則缺口', len(right_only)])

    diff_ws = wb.create_sheet('文字差異')
    diff_ws.append(['類型', '內容'])
    _apply_header(diff_ws[1])
    if added_lines:
        for item in added_lines:
            diff_ws.append(['右側新增', item])
    else:
        diff_ws.append(['右側新增', '未偵測到明顯新增內容'])
    if removed_lines:
        for item in removed_lines:
            diff_ws.append(['左側移除', item])
    else:
        diff_ws.append(['左側移除', '未偵測到明顯移除內容'])
    for row in diff_ws.iter_rows(min_row=2):
        if row[0].value == '右側新增':
            row[0].fill = GOOD_FILL
        else:
            row[0].fill = ALERT_FILL
        row[1].alignment = Alignment(wrap_text=True, vertical='top')

    issue_ws = wb.create_sheet('規則缺口')
    issue_ws.append(['差異側', '代碼', '標題', '嚴重度', '描述'])
    _apply_header(issue_ws[1])
    if left_only:
        for item in left_only:
            issue_ws.append(['僅左側', item.get('code', ''), item.get('title', ''), item.get('severity', ''), item.get('description', '')])
    if right_only:
        for item in right_only:
            issue_ws.append(['僅右側', item.get('code', ''), item.get('title', ''), item.get('severity', ''), item.get('description', '')])
    if not left_only and not right_only:
        issue_ws.append(['-', '-', '未發現規則缺口差異', '-', '-'])
    for row in issue_ws.iter_rows(min_row=2):
        row[4].alignment = Alignment(wrap_text=True, vertical='top')

    citation_ws = wb.create_sheet('引用來源')
    citation_ws.append(['文件', '來源路徑', '片段摘要'])
    _apply_header(citation_ws[1])
    if citations:
        for item in citations:
            citation_ws.append([item.get('title', ''), item.get('source_path', ''), item.get('preview', '')])
    else:
        citation_ws.append(['', '', '無'])
    for row in citation_ws.iter_rows(min_row=2):
        row[2].alignment = Alignment(wrap_text=True, vertical='top')

    for ws in wb.worksheets:
        ws.freeze_panes = 'A2'
        _fit_columns(ws)

    output = BytesIO()
    wb.save(output)
    return output.getvalue()


def build_document_compare_docx(compare_data: dict) -> bytes:
    document = WordDocument()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.start_type = WD_SECTION.NEW_PAGE

    styles = document.styles
    styles['Normal'].font.name = 'Microsoft JhengHei'
    styles['Normal'].font.size = Pt(10.5)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run(f'{COMPANY_NAME} 文件差異比對報告')
    header_run.bold = True
    header_run.font.size = Pt(9)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    _set_page_number(footer_p)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(COMPANY_NAME)
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run('ISO 文件差異與版次覆核報告')
    subrun.bold = True
    subrun.font.size = Pt(14)

    cover = document.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.add_run('文件類型：正式覆核模板 / 供品質稽核、改版審查、文件差異確認使用').italic = True

    left_doc = compare_data.get('left_document', {})
    right_doc = compare_data.get('right_document', {})
    _add_key_value_table(document, [
        ('報告編號', REPORT_CODE),
        ('產出時間', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ('左側文件', left_doc.get('title', '')),
        ('右側文件', right_doc.get('title', '')),
        ('左側來源', left_doc.get('source_path', '')),
        ('右側來源', right_doc.get('source_path', '')),
        ('文字相似度', str(compare_data.get('similarity', ''))),
        ('Prompt 版本', compare_data.get('prompt_version', '')),
        ('需人工覆核', '是' if compare_data.get('needs_human_review') else '否'),
    ])

    document.add_heading('簽核欄位', level=2)
    sign_table = document.add_table(rows=2, cols=3)
    sign_table.style = 'Table Grid'
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_headers = ['製表 / 稽核人員', '覆核', '核准']
    for idx, header_text in enumerate(sign_headers):
        sign_table.rows[0].cells[idx].text = header_text
        sign_table.rows[1].cells[idx].text = '\n\n'
    _set_table_header(sign_table.rows[0], fill_hex='4F81BD')
    document.add_paragraph()

    document.add_heading('版次結論', level=2)
    document.add_paragraph(compare_data.get('version_change_conclusion', ''))
    document.add_paragraph(compare_data.get('version_change_recommendation', ''))

    document.add_heading('比對摘要', level=2)
    summary_p = document.add_paragraph(compare_data.get('summary', ''))
    summary_p.paragraph_format.space_after = Pt(8)

    _add_list_table(document, '差異統計', ['項目', '數量'], [
        ['右側新增內容', str(len(compare_data.get('added_lines', [])))],
        ['左側移除內容', str(len(compare_data.get('removed_lines', [])))],
        ['僅左側存在的規則缺口', str(len(compare_data.get('left_only_issues', [])))],
        ['僅右側存在的規則缺口', str(len(compare_data.get('right_only_issues', [])))],
    ])

    _add_list_table(
        document,
        '文字差異',
        ['類型', '內容'],
        [[ '右側新增', item ] for item in (compare_data.get('added_lines') or ['未偵測到明顯新增內容'])] +
        [[ '左側移除', item ] for item in (compare_data.get('removed_lines') or ['未偵測到明顯移除內容'])],
    )

    issue_rows = []
    for item in compare_data.get('left_only_issues', []):
        issue_rows.append(['僅左側', item.get('code', ''), item.get('title', ''), item.get('severity', ''), item.get('description', '')])
    for item in compare_data.get('right_only_issues', []):
        issue_rows.append(['僅右側', item.get('code', ''), item.get('title', ''), item.get('severity', ''), item.get('description', '')])
    if not issue_rows:
        issue_rows.append(['-', '-', '未發現規則缺口差異', '-', '-'])
    _add_list_table(document, '規則缺口', ['差異側', '代碼', '標題', '嚴重度', '描述'], issue_rows)

    citation_rows = []
    for item in compare_data.get('citations', []):
        citation_rows.append([item.get('title', ''), item.get('source_path', ''), item.get('preview', '')])
    if not citation_rows:
        citation_rows.append(['', '', '無'])
    _add_list_table(document, '引用來源', ['文件', '來源路徑', '片段摘要'], citation_rows)

    document.add_heading('覆核說明', level=2)
    review_notes = [
        '1. 本報告為系統自動產生之差異與缺漏輔助資料，不能取代正式文件審查。',
        '2. 若判定為同文件不同版次，仍需逐條確認新增、移除與規則缺口是否符合實際修訂目的。',
        '3. 若引用來源不足或文字抽取不完整，應回到原始文件進行人工比對。',
    ]
    for note in review_notes:
        document.add_paragraph(note, style='List Bullet')

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_document_audit_docx(audit_data: dict) -> bytes:
    document = WordDocument()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"].font.size = Pt(10.5)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run(f"{COMPANY_NAME} AI 文件稽核報告")
    header_run.bold = True
    header_run.font.size = Pt(9)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    _set_page_number(footer_p)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(COMPANY_NAME)
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("AI 文件稽核正式報告")
    subrun.bold = True
    subrun.font.size = Pt(14)

    _add_key_value_table(document, [
        ("報告編號", "AA-V2-AUD-01"),
        ("產出時間", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("文件名稱", audit_data.get("document_title", "")),
        ("文件路徑", audit_data.get("document_path", "")),
        ("Prompt 版本", audit_data.get("prompt_version", "")),
        ("需人工覆核", "是" if audit_data.get("needs_human_review") else "否"),
    ])

    document.add_heading("簽核欄位", level=2)
    sign_table = document.add_table(rows=2, cols=3)
    sign_table.style = "Table Grid"
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header_text in enumerate(["製表 / 稽核人員", "覆核", "核准"]):
        sign_table.rows[0].cells[idx].text = header_text
        sign_table.rows[1].cells[idx].text = "\n\n"
    _set_table_header(sign_table.rows[0], fill_hex="4F81BD")
    document.add_paragraph()

    document.add_heading("稽核摘要", level=2)
    document.add_paragraph(audit_data.get("summary", ""))

    issue_rows = []
    for item in audit_data.get("issues", []):
        issue_rows.append([
            item.get("code", ""),
            item.get("title", ""),
            item.get("severity", ""),
            item.get("description", ""),
        ])
    if not issue_rows:
        issue_rows.append(["-", "未發現必要章節缺漏", "-", "-"])
    _add_list_table(document, "問題清單", ["代碼", "標題", "嚴重度", "描述"], issue_rows)

    evidence_rows = []
    for item in audit_data.get("insufficient_evidence", []):
        evidence_rows.append([item])
    if not evidence_rows:
        evidence_rows.append(["無"])
    _add_list_table(document, "證據不足", ["說明"], evidence_rows)

    citation_rows = []
    for item in audit_data.get("citations", []):
        citation_rows.append([
            item.get("title", ""),
            item.get("source_path", ""),
            item.get("preview", ""),
        ])
    if not citation_rows:
        citation_rows.append(["", "", "無"])
    _add_list_table(document, "引用來源", ["文件", "來源路徑", "片段摘要"], citation_rows)

    document.add_heading("覆核說明", level=2)
    for note in [
        "1. 本報告為系統規則稽核結果與引用摘要，不能取代正式文件核准流程。",
        "2. 若文件版面、簽核欄位或圖片內容未完整抽出，仍應回原始檔人工確認。",
        "3. 對高嚴重度缺口，應啟動修訂或矯正措施流程。"]:
        document.add_paragraph(note, style="List Bullet")

    output = BytesIO()
    document.save(output)
    return output.getvalue()
