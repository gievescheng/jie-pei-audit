import io
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest import mock

from docx import Document
from openpyxl import Workbook, load_workbook

import ops_data
import record_imports
import server


class FlaskOperationsTest(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.original_root = ops_data._STORAGE_ROOT
        self.original_defaults = {kind: list(meta["defaults"]) for kind, meta in ops_data.KIND_META.items()}
        ops_data.set_storage_root(Path(self.temp_dir.name))
        for kind in ops_data.KIND_META:
            ops_data.KIND_META[kind]["defaults"] = []
            ops_data.save_records(kind, [])
        server.app.config["TESTING"] = True
        self.client = server.app.test_client()

    def tearDown(self):
        for kind, defaults in self.original_defaults.items():
            ops_data.KIND_META[kind]["defaults"] = defaults
        ops_data.set_storage_root(self.original_root)
        self.temp_dir.cleanup()

    def test_nonconformance_save_delete_and_import(self):
        response = self.client.get("/api/nonconformances")
        self.assertEqual(response.status_code, 200)
        initial_count = len(response.get_json()["items"])

        create_response = self.client.post(
            "/api/nonconformances",
            json={
                "record": {
                    "date": "2026-03-01",
                    "dept": "品管課",
                    "type": "製程異常",
                    "description": "測試用不符合",
                    "severity": "輕微",
                    "responsible": "王小明",
                }
            },
        )
        self.assertEqual(create_response.status_code, 200)
        created_items = create_response.get_json()["items"]
        self.assertEqual(len(created_items), initial_count + 1)
        created_id = created_items[-1]["id"]

        doc = Document()
        doc.add_paragraph("編號： MR15-TEST-001")
        doc.add_paragraph("發現單位：管理部")
        doc.add_paragraph("不符合日期：2026/03/02")
        doc.add_paragraph("不符合事項說明：檢驗時玻璃掉落破裂。 發現者/稽核員：王小明")
        doc.add_paragraph("原因分析：作業中碰撞到 FOSB。")
        doc.add_paragraph("矯正措施：立即清場並重新教育訓練。 預定完成日期：2026/03/03")
        doc.add_paragraph("最終查核確認：■結案日期：2026/03/04")
        doc_path = Path(self.temp_dir.name) / "nc_import.docx"
        doc.save(doc_path)

        import_response = self.client.post(
            "/api/nonconformances/import",
            data={"file": (io.BytesIO(doc_path.read_bytes()), "nc_import.docx")},
            content_type="multipart/form-data",
        )
        self.assertEqual(import_response.status_code, 200)
        draft = import_response.get_json()["draft"]
        self.assertEqual(draft["dept"], "管理部")
        self.assertEqual(draft["responsible"], "王小明")
        self.assertEqual(draft["date"], "2026-03-02")

        delete_response = self.client.delete(f"/api/nonconformances/{created_id}")
        self.assertEqual(delete_response.status_code, 200)
        self.assertEqual(len(delete_response.get_json()["items"]), initial_count)

    def test_audit_plan_docx_import_and_attachments(self):
        doc = Document()
        table = doc.add_table(rows=4, cols=6)
        table.rows[0].cells[0].text = "稽核時間: 2026/04/08"
        table.rows[1].cells[0].text = "NO"
        table.rows[1].cells[1].text = "受稽核單位"
        table.rows[1].cells[3].text = "受稽核人員"
        table.rows[1].cells[4].text = "稽核員"
        table.rows[1].cells[5].text = "稽核內容"
        table.rows[2].cells[0].text = "1"
        table.rows[2].cells[1].text = "品管課"
        table.rows[2].cells[3].text = "林佑翰"
        table.rows[2].cells[4].text = "王稽核"
        table.rows[2].cells[5].text = "文件化資訊管制程序"
        doc_path = Path(self.temp_dir.name) / "audit_plan.docx"
        doc.save(doc_path)

        import_response = self.client.post(
            "/api/audit-plans/import",
            data={"file": (io.BytesIO(doc_path.read_bytes()), "audit_plan.docx")},
            content_type="multipart/form-data",
        )
        self.assertEqual(import_response.status_code, 200)
        records = import_response.get_json()["records"]
        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]["dept"], "品管課")
        self.assertEqual(records[0]["scope"], "文件化資訊管制程序")
        self.assertEqual(records[0]["scheduledDate"], "2026-04-08")

        save_response = self.client.post("/api/audit-plans", json={"records": records})
        self.assertEqual(save_response.status_code, 200)
        saved_id = save_response.get_json()["saved"][0]["id"]

        attachment_response = self.client.get(f"/api/audit-plans/{saved_id}/attachments")
        self.assertEqual(attachment_response.status_code, 200)
        attachments = attachment_response.get_json()["attachments"]
        self.assertEqual(len(attachments), 1)
        self.assertTrue(attachments[0]["text_previewable"])

        preview_response = self.client.get(attachments[0]["preview_text_url"])
        self.assertEqual(preview_response.status_code, 200)
        self.assertIn("audit_plan.docx", preview_response.get_data(as_text=True))

        delete_response = self.client.delete(f"/api/audit-plans/{saved_id}")
        self.assertEqual(delete_response.status_code, 200)

    def test_environment_import_filter_and_delete_range(self):
        wb = Workbook()
        ws = wb.active
        ws.title = "Clean_Long"
        ws.append(["Model:9303"])
        ws.append(["Count Mode:Σ"])
        ws.append(["Number of Samples: 2"])
        ws.append(["Channel", "", "Min", "Max"])
        ws.append(["1(0.3)"])
        ws.append(["2(0.5)"])
        ws.append(["3(5.0)"])
        ws.append(["Date", "Point", "DateTime", "Record", "Ch1(um)", "Ch1_count", "Ch2(um)", "Ch2_count", "Ch3(um)", "Ch3_count", "SampleTime", "HoldTime"])
        ws.append(["2026/03/01", 1, "2026/03/01 08:00:00", 1, 0.3, 100, 0.5, 20, 5.0, 3, "00:01:00", "00:00:05"])
        ws.append(["2026/03/05", 2, "2026/03/05 08:00:00", 2, 0.3, 1500, 0.5, 900, 5.0, 40, "00:01:00", "00:00:05"])
        ws.append(["2026/03/06 09:10:11", "", "", "0.3", "", 61, "", "0.5", 4, "", "", "5.0"])
        sample_path = Path(self.temp_dir.name) / "environment_clean_long.xlsx"
        wb.save(sample_path)

        import_response = self.client.post(
            "/api/environment-records/import",
            data={"file": (io.BytesIO(sample_path.read_bytes()), "environment_clean_long.xlsx")},
            content_type="multipart/form-data",
        )
        self.assertEqual(import_response.status_code, 200)
        records = import_response.get_json()["records"]
        self.assertEqual(len(records), 3)
        self.assertEqual(records[0]["location"], "粒子計數點 1 (08:00:00)")
        self.assertEqual(records[2]["location"], "粒子計數點 1 (09:10:11)")
        self.assertEqual(records[2]["point"], "1")
        self.assertEqual(records[2]["particles03"], 61)
        self.assertEqual(records[2]["particles05"], 4)

        save_response = self.client.post(
            "/api/environment-records",
            json={"records": records, "replace_source_file": records[0]["source_file"]},
        )
        self.assertEqual(save_response.status_code, 200)
        self.assertEqual(len(save_response.get_json()["items"]), 3)

        second_save_response = self.client.post(
            "/api/environment-records",
            json={"records": records, "replace_source_file": records[0]["source_file"]},
        )
        self.assertEqual(second_save_response.status_code, 200)
        self.assertEqual(len(second_save_response.get_json()["items"]), 3)

        filter_response = self.client.get("/api/environment-records?start=2026-03-02&end=2026-03-31")
        self.assertEqual(filter_response.status_code, 200)
        payload = filter_response.get_json()
        self.assertEqual(len(payload["items"]), 2)
        self.assertEqual(payload["summary"]["failed"], 1)

        delete_response = self.client.post(
            "/api/environment-records/delete-range",
            json={"start": "2026-03-01", "end": "2026-03-03"},
        )
        self.assertEqual(delete_response.status_code, 200)
        deleted_payload = delete_response.get_json()
        self.assertEqual(len(deleted_payload["items"]), 2)
        self.assertEqual(deleted_payload["removed_count"], 1)

    def test_environment_template_style_import(self):
        wb = Workbook()
        ws = wb.active
        ws.title = "環境監控資料"
        ws.append(["Date", "DateTime", "Point", "Location", "0.3um", "0.5um", "5.0um", "Temp", "Humidity", "Pressure", "Operator", "Result"])
        ws.append(["日期格式", "量測時間", "1~14", "可留白", "0.3", "0.5", "5.0", "可留白", "可留白", "可留白", "記錄者", "可留白"])
        ws.append(["2026-03-10", "2026-03-10 08:30:00", "3", "", 120, 15, 2, "", "", "", "王小明", ""])
        ws.append(["2026-03-10", "2026-03-10 09:00:00", "", "A區溫濕度測點", "", "", "", 22.4, 46.0, 11.5, "陳小華", ""])
        guide = wb.create_sheet("填寫說明")
        guide["A1"] = "說明頁"
        sample_path = Path(self.temp_dir.name) / "environment_template.xlsx"
        wb.save(sample_path)

        import_response = self.client.post(
            "/api/environment-records/import",
            data={"file": (io.BytesIO(sample_path.read_bytes()), "environment_template.xlsx")},
            content_type="multipart/form-data",
        )
        self.assertEqual(import_response.status_code, 200)
        records = import_response.get_json()["records"]
        self.assertEqual(len(records), 2)
        self.assertEqual(records[0]["point"], "3")
        self.assertEqual(records[0]["location"], "粒子計數點 3")
        self.assertEqual(records[0]["particles03"], 120)
        self.assertEqual(records[0]["particles05"], 15)
        self.assertEqual(records[0]["particles5"], 2)
        self.assertEqual(records[1]["temp"], 22.4)

    def test_record_engine_catalog_suggest_and_generate(self):
        catalog_response = self.client.get("/api/record-engine/catalog")
        self.assertEqual(catalog_response.status_code, 200)
        catalog = catalog_response.get_json()["templates"]
        codes = {item["code"] for item in catalog}
        self.assertIn("env_monthly_pack", codes)
        self.assertIn("shipping_pack", codes)
        self.assertIn("cip_152", codes)

        suggest_response = self.client.post(
            "/api/record-engine/suggest",
            json={
                "prompt": "我要產生出貨流程相關紀錄",
                "context": {
                    "shipment_order_count": 2,
                    "prod_count": 1,
                    "quality_count": 1,
                    "env_count": 1,
                    "nonconformance_count": 1,
                },
            },
        )
        self.assertEqual(suggest_response.status_code, 200)
        suggested = suggest_response.get_json()["templates"]
        self.assertGreater(len(suggested), 0)
        self.assertEqual(suggested[0]["code"], "shipping_pack")
        shipping_pack = next(item for item in catalog if item["code"] == "shipping_pack")
        self.assertEqual(shipping_pack["included_templates"], ["14.3 出貨單", "12.5 出貨檢查紀錄表", "14.5 出貨檢查表"])
        env_monthly_pack = next(item for item in catalog if item["code"] == "env_monthly_pack")
        self.assertEqual(env_monthly_pack["included_templates"], ["6.1 環境監控記錄表"])
        cip_pack = next(item for item in catalog if item["code"] == "cip_pack")
        self.assertEqual(cip_pack["included_templates"], ["15.2 製程缺陷追蹤改善(CIP)紀錄表"])

        env_suggest_response = self.client.post(
            "/api/record-engine/suggest",
            json={
                "prompt": "請幫我整理環境監控月報",
                "context": {
                    "env_count": 5,
                },
            },
        )
        self.assertEqual(env_suggest_response.status_code, 200)
        env_suggested = env_suggest_response.get_json()["templates"]
        self.assertGreater(len(env_suggested), 0)
        self.assertEqual(env_suggested[0]["code"], "env_monthly_pack")

        material_response = self.client.post(
            "/api/record-engine/generate",
            json={
                "template_code": "material_request_112",
                "shipment_request": {
                    "order_no": "4515994888",
                    "date": "2026-03-19",
                    "department": "資材課",
                    "requester": "測試員",
                    "product_name": "RECYCLE GLASS NEG ABC-1 (JEPE)",
                    "spec": "12吋",
                    "quantity": 200,
                    "unit": "片",
                    "remark": "測試用領料單",
                    "batch_display": "4515994888/JPAN111001",
                    "selected_lots": ["JPAN111001"],
                },
            },
        )
        self.assertEqual(material_response.status_code, 200)

    def test_record_engine_precheck(self):
        incomplete_response = self.client.post(
            "/api/record-engine/precheck",
            json={
                "template_code": "shipment_order_143",
                "shipment_request": {
                    "order_no": "4515994888",
                    "date": "",
                    "department": "",
                    "requester": "",
                    "product_name": "",
                    "quantity": "",
                },
            },
        )
        self.assertEqual(incomplete_response.status_code, 200)
        incomplete_result = incomplete_response.get_json()["result"]
        self.assertFalse(incomplete_result["ready"])
        self.assertIn("出貨日期", incomplete_result["missing_items"])
        self.assertIn("申請部門", incomplete_result["missing_items"])

        complete_response = self.client.post(
            "/api/record-engine/precheck",
            json={
                "template_code": "shipping_pack",
                "prod_records": [{"lot": "0318-1", "customer": "C001"}],
                "quality_records": [{"materialName": "WAFER", "batchNo": "B-01"}],
                "shipment_request": {
                    "order_no": "4515994888",
                    "date": "2026-03-19",
                    "department": "資材課",
                    "requester": "測試員",
                    "product_name": "RECYCLE GLASS NEG ABC-1 (JEPE)",
                    "quantity": "200",
                },
            },
        )
        self.assertEqual(complete_response.status_code, 200)
        complete_result = complete_response.get_json()["result"]
        self.assertTrue(complete_result["ready"])
        self.assertEqual(len(complete_result["included_templates"]), 3)

        cip_pack_response = self.client.post(
            "/api/record-engine/precheck",
            json={
                "template_code": "cip_pack",
                "nonconformance": {
                    "id": "NC-TEST-001",
                    "date": "2026-03-01",
                    "dept": "品管課",
                    "description": "玻璃破片流出",
                    "responsible": "王小明",
                },
            },
        )
        self.assertEqual(cip_pack_response.status_code, 200)
        cip_pack_result = cip_pack_response.get_json()["result"]
        self.assertTrue(cip_pack_result["ready"])
        self.assertEqual(len(cip_pack_result["included_templates"]), 1)
        self.assertIn("嚴重度", "\n".join(cip_pack_result["warnings"]))

        env_pack_response = self.client.post(
            "/api/record-engine/precheck",
            json={
                "template_code": "env_monthly_pack",
                "env_records": [
                    {"date": "2025-11-14", "dateTime": "2025-11-14 11:22:50", "point": "1", "location": "粒子計數點 1", "particles03": 16, "particles05": 10, "particles5": 3, "result": "合格"},
                    {"date": "2025-11-15", "dateTime": "2025-11-15 09:10:11", "point": "2", "location": "粒子計數點 2", "particles03": 1800, "particles05": 900, "particles5": 40, "result": "不合格"},
                ],
            },
        )
        self.assertEqual(env_pack_response.status_code, 200)
        env_pack_result = env_pack_response.get_json()["result"]
        self.assertTrue(env_pack_result["ready"])
        self.assertEqual(len(env_pack_result["included_templates"]), 1)
        self.assertIn("月報摘要", "\n".join(env_pack_result["warnings"]))

    def test_production_record_upload_import(self):
        wb = Workbook()
        ws = wb.active
        ws.append(["日期", "站點", "客戶/產品", "Wafer Boat Lot", "投入", "備註", "良品", "不良", "良率", "不良原因", "作業員", "補充"])
        ws.append(["2026/03/20", "OCR", "C001/待洗玻璃", "0320-1", 30, "", 28, 2, "93.3%", "刮傷, 髒污", "王小明", "夜班"])
        sample_path = Path(self.temp_dir.name) / "production_upload.xlsx"
        wb.save(sample_path)

        response = self.client.post(
            "/api/production-records/import",
            data={"file": (io.BytesIO(sample_path.read_bytes()), "production_upload.xlsx")},
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        self.assertEqual(payload["source_file"], "production_upload.xlsx")
        self.assertEqual(len(payload["records"]), 1)
        self.assertEqual(payload["records"][0]["lot"], "0320-1")
        self.assertEqual(payload["records"][0]["customer"], "C001")
        self.assertEqual(payload["records"][0]["product"], "待洗玻璃")

    def test_quality_record_upload_import(self):
        wb = Workbook()
        ws = wb.active
        for _ in range(4):
            ws.append([])
        ws.append(["材料名稱", "批號", "數量", "規格", "檢驗數量", "PH", "比重", "RI", "旋光值", "結果", "備註"])
        ws.append(["IPA", "B20260320", "200kg", "電子級", "3", "7.0", "0.98", "1.33", "0.1", "OK", "正常"])
        sample_path = Path(self.temp_dir.name) / "quality_upload.xlsx"
        wb.save(sample_path)

        response = self.client.post(
            "/api/quality-records/import",
            data={"file": (io.BytesIO(sample_path.read_bytes()), "quality_upload.xlsx")},
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        self.assertEqual(payload["source_file"], "quality_upload.xlsx")
        self.assertEqual(len(payload["records"]), 1)
        self.assertEqual(payload["records"][0]["materialName"], "IPA")
        self.assertEqual(payload["records"][0]["batchNo"], "B20260320")

    def test_record_engine_generate_outputs(self):
        material_response = self.client.post(
            "/api/record-engine/generate",
            json={
                "template_code": "material_request_112",
                "shipment_request": {
                    "order_no": "4515994888",
                    "date": "2026-03-19",
                    "department": "資材課",
                    "requester": "測試員",
                    "product_name": "RECYCLE GLASS NEG ABC-1 (JEPE)",
                    "spec": "12吋",
                    "quantity": 200,
                    "unit": "片",
                    "remark": "測試用領料單",
                    "batch_display": "4515994888/JPAN111001",
                    "selected_lots": ["JPAN111001"],
                },
            },
        )
        self.assertEqual(material_response.status_code, 200)
        wb = load_workbook(io.BytesIO(material_response.data))
        ws = wb.active
        self.assertEqual(ws["B3"].value, "資材課")
        self.assertEqual(ws["B5"].value, "4515994888/JPAN111001")
        wb.close()

        cip_response = self.client.post(
            "/api/record-engine/generate",
            json={
                "template_code": "cip_152",
                "nonconformance": {
                    "id": "NC-TEST-001",
                    "date": "2026-03-01",
                    "dept": "品管課",
                    "type": "製程異常",
                    "description": "玻璃破片流出",
                    "rootCause": "作業中碰撞到 FOSB",
                    "correctiveAction": "重新教育訓練並加嚴檢查",
                    "responsible": "王小明",
                    "dueDate": "2026-03-10",
                    "status": "進行中",
                },
            },
        )
        self.assertEqual(cip_response.status_code, 200)
        wb = load_workbook(io.BytesIO(cip_response.data))
        ws = wb.active
        self.assertEqual(ws["A3"].value, "不符合類型")
        self.assertEqual(ws["B3"].value, "製程異常")
        self.assertEqual(ws["A8"].value, "問題描述")
        self.assertEqual(ws["B8"].value, "玻璃破片流出")
        self.assertEqual(ws["A12"].value, "到期日")
        self.assertEqual(ws["B12"].value, "2026-03-10")
        self.assertEqual(ws["A13"].value, "狀態")
        self.assertEqual(ws["B13"].value, "進行中")
        self.assertIn("改善追蹤", wb.sheetnames)
        follow_ws = wb["改善追蹤"]
        self.assertEqual(follow_ws["A2"].value, "問題描述")
        wb.close()

        shipping_pack_response = self.client.post(
            "/api/record-engine/generate",
            json={
                "template_code": "shipping_pack",
                "shipment_request": {
                    "order_no": "4515994888",
                    "date": "2026-03-19",
                    "department": "資材課",
                    "requester": "測試員",
                    "product_name": "RECYCLE GLASS NEG ABC-1 (JEPE)",
                    "spec": "12吋",
                    "quantity": 200,
                    "unit": "片",
                    "remark": "測試用出貨流程",
                    "batch_display": "4515994888/JPAN111001",
                    "selected_lots": ["JPAN111001"],
                },
            },
        )
        self.assertEqual(shipping_pack_response.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(shipping_pack_response.data)) as archive:
            names = set(archive.namelist())
        self.assertEqual(len(names), 3)
        self.assertTrue(any("出貨單" in name for name in names))
        self.assertTrue(any("出貨檢查紀錄表" in name for name in names))

        cip_pack_response = self.client.post(
            "/api/record-engine/generate",
            json={
                "template_code": "cip_pack",
                "nonconformance": {
                    "id": "NC-TEST-001",
                    "date": "2026-03-01",
                    "dept": "品管課",
                    "type": "製程異常",
                    "description": "玻璃破片流出",
                    "rootCause": "作業中碰撞到 FOSB",
                    "correctiveAction": "重新教育訓練並加嚴檢查",
                    "responsible": "王小明",
                    "dueDate": "2026-03-10",
                    "status": "進行中",
                },
            },
        )
        self.assertEqual(cip_pack_response.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(cip_pack_response.data)) as archive:
            names = set(archive.namelist())
            summary_name = next(name for name in names if "不符合來源摘要" in name)
            summary_wb = load_workbook(io.BytesIO(archive.read(summary_name)))
        self.assertEqual(len(names), 2)
        self.assertTrue(any("CIP紀錄表" in name for name in names))
        self.assertTrue(any("不符合來源摘要" in name for name in names))
        summary_ws = summary_wb.active
        self.assertEqual(summary_ws["A8"].value, "問題描述")
        self.assertEqual(summary_ws["A14"].value, "結案日期")
        summary_wb.close()

        env_pack_response = self.client.post(
            "/api/record-engine/generate",
            json={
                "template_code": "env_monthly_pack",
                "env_records": [
                    {"date": "2025-11-14", "dateTime": "2025-11-14 11:22:50", "point": "1", "location": "粒子計數點 1", "particles03": 16, "particles05": 10, "particles5": 3, "operator": "王小明", "result": "合格"},
                    {"date": "2025-11-15", "dateTime": "2025-11-15 09:10:11", "point": "2", "location": "粒子計數點 2", "particles03": 1800, "particles05": 900, "particles5": 40, "operator": "陳小華", "result": "不合格"},
                ],
            },
        )
        self.assertEqual(env_pack_response.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(env_pack_response.data)) as archive:
            names = set(archive.namelist())
            summary_name = next(name for name in names if "環境監控月報" in name)
            detail_name = next(name for name in names if "環境監控記錄表" in name)
            summary_wb = load_workbook(io.BytesIO(archive.read(summary_name)))
            detail_wb = load_workbook(io.BytesIO(archive.read(detail_name)))
        self.assertEqual(len(names), 2)
        # 格式化後封面為第一張工作表，月報摘要為第二張（index 1）
        self.assertIn("封面", summary_wb.sheetnames)
        summary_ws = summary_wb.worksheets[1]
        self.assertEqual(summary_ws.title, "月報摘要")
        self.assertEqual(summary_ws["A1"].value, "6 工作環境監控月報摘要")
        self.assertIn("日別統計", summary_wb.sheetnames)
        self.assertIn("點位統計", summary_wb.sheetnames)
        # 環境監控記錄表：A1=標題「日期」，A2=第一筆資料日期
        self.assertEqual(detail_wb.active["A1"].value, "日期")
        self.assertIsNotNone(detail_wb.active["A2"].value)
        summary_wb.close()
        detail_wb.close()

    def test_existing_production_and_quality_record_read(self):
        prod_wb = Workbook()
        prod_ws = prod_wb.active
        prod_ws.title = "生產日報"
        prod_ws.append(["潔沛企業有限公司"])
        prod_ws.append(["生產日報表"])
        prod_ws.append(["年度", "115", "表單編號", "MR11-05", "", "", "文件版次", "B", "訂單編號", "", "建立日期", "2026.03.18"])
        prod_ws.append(["日期", "設施站點", "客戶代號/產品名稱", "Wafer Boat Lot", "投入數", "篩選不用洗", "良品數", "不良品數", "良率", "不良原因/數量", "生產人員", "備註"])
        prod_ws.append(["2026.03.18", "OCR", "C001/待洗玻璃", "0318-1", "25", "0", "25", "0", "1", "", "楊", ""])
        prod_path = Path(self.temp_dir.name) / "prod.xlsx"
        prod_wb.save(prod_path)

        quality_wb = Workbook()
        quality_ws = quality_wb.active
        quality_ws.title = "品質"
        quality_ws.append(["潔沛企業有限公司"])
        quality_ws.append(["編訂部門:品管課"])
        quality_ws.append(["文件編號:MR12-01"])
        quality_ws.append(["注意事項"])
        quality_ws.append(["原料名稱", "原料批號", "原料數量", "規格", "品檢數量", "PH值檢驗", "比重值檢驗", "RI值檢驗", "旋光度檢驗", "外觀", "備註"])
        quality_ws.append(["WAFER", "B-01", "675", "片", "27盒", "NA", "NA", "NA", "NA", "OK", "外觀無破損"])
        quality_path = Path(self.temp_dir.name) / "quality.xlsx"
        quality_wb.save(quality_path)

        with mock.patch.object(record_imports, "load_existing_production_records", return_value=(record_imports.parse_production_record_file(prod_path), str(prod_path))):
            response = self.client.get("/api/production-records/read-existing")
            self.assertEqual(response.status_code, 200)
            payload = response.get_json()
            self.assertEqual(len(payload["records"]), 1)
            self.assertEqual(payload["records"][0]["lot"], "0318-1")
            self.assertEqual(payload["records"][0]["customer"], "C001")
            self.assertEqual(payload["records"][0]["yieldRate"], 100.0)

        with mock.patch.object(record_imports, "load_existing_quality_records", return_value=(record_imports.parse_quality_record_file(quality_path), str(quality_path))):
            response = self.client.get("/api/quality-records/read-existing")
            self.assertEqual(response.status_code, 200)
            payload = response.get_json()
            self.assertEqual(len(payload["records"]), 1)
            self.assertEqual(payload["records"][0]["materialName"], "WAFER")
            self.assertEqual(payload["records"][0]["result"], "PASS")


if __name__ == "__main__":
    unittest.main()
